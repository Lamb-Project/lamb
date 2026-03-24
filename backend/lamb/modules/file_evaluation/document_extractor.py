"""
Document text extraction service.
Extracts text from PDF, DOCX, and plain-text/code files.

Ported from LAMBA with LAMB-style logging.
"""
import os
from typing import Optional
from lamb.logging_config import get_logger

logger = get_logger(__name__, component="FILE_EVAL")

try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("pypdf not installed – PDF extraction disabled")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed – DOCX extraction disabled")


_TEXT_EXTENSIONS = frozenset([
    '.txt', '.md', '.py', '.java', '.cpp', '.c', '.js',
    '.html', '.css', '.json', '.xml', '.ts', '.rs', '.go',
    '.rb', '.sh', '.yml', '.yaml', '.toml', '.ini', '.cfg',
])


class DocumentExtractor:
    """Extract text from various document formats."""

    @staticmethod
    def extract_text_from_file(file_path: str) -> Optional[str]:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None

        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        try:
            if ext == '.pdf':
                return DocumentExtractor._extract_from_pdf(file_path)
            elif ext in ('.docx', '.doc'):
                return DocumentExtractor._extract_from_docx(file_path)
            elif ext in _TEXT_EXTENSIONS:
                return DocumentExtractor._extract_from_text(file_path)
            else:
                logger.warning(f"Unsupported file format: {ext}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None

    @staticmethod
    def _extract_from_pdf(file_path: str) -> Optional[str]:
        if not PDF_AVAILABLE:
            logger.error("pypdf not available")
            return None
        parts = []
        with open(file_path, 'rb') as fh:
            for page in PdfReader(fh).pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
        full = '\n'.join(parts)
        return full.strip() or None

    @staticmethod
    def _extract_from_docx(file_path: str) -> Optional[str]:
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available")
            return None
        doc = docx.Document(file_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        full = '\n'.join(parts)
        return full.strip() or None

    @staticmethod
    def _extract_from_text(file_path: str) -> Optional[str]:
        try:
            with open(file_path, 'r', encoding='utf-8') as fh:
                text = fh.read()
            return text.strip() or None
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as fh:
                text = fh.read()
            return text.strip() or None

    @staticmethod
    def get_text_preview(text: str, max_length: int = 500) -> str:
        if not text:
            return ""
        if len(text) <= max_length:
            return text
        return text[:max_length] + "..."
